from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Title Page
title = doc.add_heading('Inventory Management System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Complete System Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(16)

date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc = [
    '1. Introduction',
    '2. Statement of the Problem (SOP)',
    '3. Objectives (Solutions)',
    '4. Scope',
    '5. Limitations',
    '6. System Features'
]
for item in toc:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. INTRODUCTION
doc.add_heading('1. INTRODUCTION', 1)

doc.add_paragraph(
    'The Inventory Management System is a comprehensive Django-based web application designed to streamline '
    'inventory operations, user management, attendance tracking, and supplier collaboration. This enterprise-grade '
    'solution provides role-based access control with four distinct user types: Super Admin, Admin, Staff, and Supplier.'
)

doc.add_paragraph(
    'Built with security as the top priority, the system implements multiple layers of protection including SQL '
    'injection prevention, CSRF protection, XSS protection, and comprehensive audit logging. The application supports '
    'multi-user environments with granular permission management and complete activity tracking for compliance and security.'
)

doc.add_paragraph(
    'The system is designed for businesses that require precise inventory control, production tracking, supplier '
    'management, and employee attendance monitoring. It features automated workflows, real-time stock tracking, '
    'expiration management, and QR code-based purchase order processing.'
)

doc.add_heading('1.1 Technology Stack', 2)
doc.add_paragraph('• Backend Framework: Django 5.1.3')
doc.add_paragraph('• Database: MySQL 5.7+')
doc.add_paragraph('• Programming Language: Python 3.8+')
doc.add_paragraph('• Timezone: Asia/Manila')
doc.add_paragraph('• Security: Django ORM, CSRF Protection, HSTS, XSS Protection')

doc.add_page_break()

# 2. STATEMENT OF THE PROBLEM (SOP)
doc.add_heading('2. STATEMENT OF THE PROBLEM (SOP)', 1)

doc.add_paragraph(
    'Traditional inventory management systems face numerous challenges that lead to inefficiencies, losses, '
    'and security vulnerabilities. This section outlines the key problems that businesses commonly encounter '
    'in inventory management, user administration, production tracking, and supplier coordination.'
)

doc.add_heading('2.1 Inventory Tracking Problems', 2)

doc.add_heading('Problem 1: Inaccurate Stock Levels', 3)
doc.add_paragraph(
    'Manual inventory tracking leads to discrepancies between actual and recorded stock levels. '
    'Businesses struggle with stock-outs or overstocking due to lack of real-time visibility into inventory quantities.'
)

doc.add_heading('Problem 2: Expired and Wasted Products', 3)
doc.add_paragraph(
    'Perishable items expire before use due to poor tracking of expiration dates and lack of FIFO (First In First Out) '
    'implementation. This results in significant financial losses and waste, especially in food and pharmaceutical industries.'
)

doc.add_heading('Problem 3: No Low Stock Alerts', 3)
doc.add_paragraph(
    'Without automated reorder alerts, businesses run out of critical items unexpectedly, causing production delays '
    'and lost sales opportunities. Manual monitoring is time-consuming and prone to human error.'
)

doc.add_heading('Problem 4: Difficulty Tracking Stock Movements', 3)
doc.add_paragraph(
    'Businesses cannot easily trace where inventory went (consumed, transferred, spoiled) or who performed the action. '
    'This lack of traceability makes it difficult to identify theft, errors, or process inefficiencies.'
)

doc.add_heading('Problem 5: Manual Item Code Generation', 3)
doc.add_paragraph(
    'Creating unique item codes manually is time-consuming and leads to duplicate codes or inconsistent naming conventions. '
    'This causes confusion and errors in inventory records.'
)

doc.add_heading('2.2 User Management and Security Problems', 2)

doc.add_heading('Problem 6: Unauthorized Access to Sensitive Data', 3)
doc.add_paragraph(
    'Without proper role-based access control, employees can access and modify data beyond their authorization level. '
    'This poses security risks and compliance issues.'
)

doc.add_heading('Problem 7: No Audit Trail', 3)
doc.add_paragraph(
    'When problems occur (missing inventory, unauthorized changes), there is no way to track who did what and when. '
    'This makes accountability and troubleshooting nearly impossible.'
)

doc.add_heading('Problem 8: Complex Permission Management', 3)
doc.add_paragraph(
    'Assigning and managing granular permissions for different staff members is cumbersome in traditional systems. '
    'Admins struggle to control who can view, edit, or delete specific types of data.'
)

doc.add_heading('Problem 9: SQL Injection and Security Vulnerabilities', 3)
doc.add_paragraph(
    'Many inventory systems are vulnerable to SQL injection attacks, data breaches, and cross-site scripting (XSS) attacks. '
    'This puts sensitive business data at risk.'
)

doc.add_heading('2.3 Production Management Problems', 2)

doc.add_heading('Problem 10: Manual Ingredient Calculation', 3)
doc.add_paragraph(
    'Production staff must manually calculate ingredient quantities needed for each batch, leading to errors, '
    'over-consumption, or under-consumption of raw materials.'
)

doc.add_heading('Problem 11: No Recipe Standardization', 3)
doc.add_paragraph(
    'Without standardized recipes, product quality varies between batches. Different staff members use different '
    'ingredient quantities, affecting consistency and cost control.'
)

doc.add_heading('Problem 12: Ingredient Waste Due to Loss Factors', 3)
doc.add_paragraph(
    'Recipes do not account for loss factors (spillage, trimming, evaporation), causing actual ingredient usage '
    'to exceed planned amounts. This leads to unexpected stock depletion.'
)

doc.add_heading('Problem 13: No Production History', 3)
doc.add_paragraph(
    'Businesses cannot track what was produced, when, by whom, and using which ingredients. This makes it difficult '
    'to analyze production efficiency or trace quality issues.'
)

doc.add_heading('2.4 Supplier and Purchase Order Problems', 2)

doc.add_heading('Problem 14: Manual Purchase Order Processing', 3)
doc.add_paragraph(
    'Creating purchase orders manually through phone calls, emails, or paper forms is slow and error-prone. '
    'Order details are often miscommunicated or lost.'
)

doc.add_heading('Problem 15: No Supplier Collaboration Platform', 3)
doc.add_paragraph(
    'Suppliers have no direct access to view orders, confirm availability, or update pricing. All communication '
    'happens through intermediaries, causing delays and miscommunication.'
)

doc.add_heading('Problem 16: Difficulty Tracking Order Status', 3)
doc.add_paragraph(
    'Businesses cannot easily track whether orders are pending, approved, shipped, or received. This leads to '
    'confusion, duplicate orders, or missed deliveries.'
)

doc.add_heading('Problem 17: Manual Receiving Process', 3)
doc.add_paragraph(
    'When deliveries arrive, staff must manually verify items, create stock lots, and update inventory records. '
    'This is time-consuming and prone to data entry errors.'
)

doc.add_heading('Problem 18: No QR Code or Barcode Tracking', 3)
doc.add_paragraph(
    'Without QR codes or barcodes, tracking physical orders and matching them to digital records is difficult. '
    'This causes receiving errors and inventory discrepancies.'
)

doc.add_heading('2.5 Employee Attendance Problems', 2)

doc.add_heading('Problem 19: Manual Attendance Recording', 3)
doc.add_paragraph(
    'Paper-based or manual attendance systems are easily manipulated (buddy punching) and difficult to verify. '
    'Managers spend excessive time compiling and verifying attendance records.'
)

doc.add_heading('Problem 20: No Shift-Based Tracking', 3)
doc.add_paragraph(
    'Businesses with AM/PM shifts cannot accurately track when employees clock in and out for each shift. '
    'This makes payroll calculation and shift management complicated.'
)

doc.add_heading('Problem 21: No Attendance History Visibility', 3)
doc.add_paragraph(
    'Employees and managers cannot easily view historical attendance records or identify patterns of absences or tardiness. '
    'This makes performance reviews and scheduling difficult.'
)

doc.add_heading('2.6 Reporting and Decision-Making Problems', 2)

doc.add_heading('Problem 22: Lack of Real-Time Dashboards', 3)
doc.add_paragraph(
    'Managers cannot get instant visibility into key metrics like total inventory value, low stock items, or production output. '
    'They must manually compile reports, which are often outdated by the time they are ready.'
)

doc.add_heading('Problem 23: No Trend Analysis', 3)
doc.add_paragraph(
    'Without historical data visualization, businesses cannot identify trends in inventory value, consumption patterns, '
    'or production efficiency. This limits strategic planning and forecasting.'
)

doc.add_heading('Problem 24: Difficult to Generate Reports', 3)
doc.add_paragraph(
    'Creating reports for stock levels, expiring items, or production summaries requires manual data extraction and formatting. '
    'This is time-consuming and limits data-driven decision making.'
)

doc.add_page_break()

# 3. OBJECTIVES (SOLUTIONS)
doc.add_heading('3. OBJECTIVES (SOLUTIONS)', 1)

doc.add_paragraph(
    'The Inventory Management System addresses all the problems outlined in the Statement of the Problem section '
    'through comprehensive, automated solutions. Each objective directly solves one or more identified problems.'
)

doc.add_heading('3.1 Solutions to Inventory Tracking Problems', 2)

doc.add_heading('Solution 1: Real-Time Stock Level Tracking', 3)
doc.add_paragraph(
    '✓ Implement automatic stock calculations that update in real-time with every receive, consume, or production transaction'
)
doc.add_paragraph('✓ Provide instant visibility into current stock quantities for all items')
doc.add_paragraph('✓ Eliminate discrepancies through automated tracking instead of manual counting')
doc.add_paragraph('✓ Display stock status (In Stock, Low Stock, Out of Stock) on dashboards')

doc.add_heading('Solution 2: Expiration Date Management and FIFO', 3)
doc.add_paragraph('✓ Track expiration dates for all perishable items at the lot level')
doc.add_paragraph('✓ Implement FIFO (First In First Out) automatic consumption to use oldest stock first')
doc.add_paragraph('✓ Generate alerts for items expiring within 7 days')
doc.add_paragraph('✓ Identify expired items automatically')
doc.add_paragraph('✓ Provide expiration tracker dashboard for easy monitoring')

doc.add_heading('Solution 3: Automated Low Stock Alerts', 3)
doc.add_paragraph('✓ Set reorder levels for each item')
doc.add_paragraph('✓ Automatically detect when stock falls below reorder level')
doc.add_paragraph('✓ Display low stock and out-of-stock counts on dashboard')
doc.add_paragraph('✓ Enable proactive reordering to prevent stock-outs')

doc.add_heading('Solution 4: Complete Stock Movement Tracking', 3)
doc.add_paragraph('✓ Record all stock movements (Receive, Consume, Produce, Adjust, Transfer, Spoilage)')
doc.add_paragraph('✓ Track who performed each movement and when')
doc.add_paragraph('✓ Maintain reference numbers and reasons for each movement')
doc.add_paragraph('✓ Provide searchable movement history for traceability')
doc.add_paragraph('✓ Enable identification of theft, errors, or inefficiencies')

doc.add_heading('Solution 5: Auto-Generated Item Codes', 3)
doc.add_paragraph('✓ Automatically generate unique item codes in format YYYYMM0001')
doc.add_paragraph('✓ Ensure no duplicate codes')
doc.add_paragraph('✓ Maintain consistent naming convention')
doc.add_paragraph('✓ Save time and eliminate manual code creation errors')

doc.add_heading('3.2 Solutions to User Management and Security Problems', 2)

doc.add_heading('Solution 6: Role-Based Access Control (RBAC)', 3)
doc.add_paragraph('✓ Implement 4 user roles: Super Admin, Admin, Staff, Supplier')
doc.add_paragraph('✓ Restrict access based on role (e.g., Staff cannot manage users)')
doc.add_paragraph('✓ Enforce role hierarchy (Admin cannot create Super Admin)')
doc.add_paragraph('✓ Protect sensitive data from unauthorized access')

doc.add_heading('Solution 7: Comprehensive Audit Logging', 3)
doc.add_paragraph('✓ Log every user action (Create, Read, Update, Delete, Login, Logout)')
doc.add_paragraph('✓ Record IP address, user agent, and timestamp for each action')
doc.add_paragraph('✓ Maintain immutable audit trail for compliance')
doc.add_paragraph('✓ Enable troubleshooting and accountability')
doc.add_paragraph('✓ Provide searchable and filterable audit logs')

doc.add_heading('Solution 8: Granular Permission Management', 3)
doc.add_paragraph('✓ Implement 10 permission types (Inventory Read/Write/Delete, User Read/Write/Delete, Reports, Settings)')
doc.add_paragraph('✓ Allow admins to assign specific permissions to staff')
doc.add_paragraph('✓ Support bulk permission updates')
doc.add_paragraph('✓ Enable permission expiration dates')
doc.add_paragraph('✓ Simplify permission management through user-friendly interface')

doc.add_heading('Solution 9: Multi-Layer Security Protection', 3)
doc.add_paragraph('✓ Prevent SQL injection through Django ORM (no raw SQL queries)')
doc.add_paragraph('✓ Implement CSRF protection on all forms')
doc.add_paragraph('✓ Sanitize and validate all user inputs to prevent XSS attacks')
doc.add_paragraph('✓ Apply secure headers (HSTS, X-Frame-Options, Content-Type-Nosniff)')
doc.add_paragraph('✓ Enforce strong password validation')

doc.add_heading('3.3 Solutions to Production Management Problems', 2)

doc.add_heading('Solution 10: Automated Ingredient Calculation', 3)
doc.add_paragraph('✓ Calculate ingredient quantities automatically based on production quantity and recipe')
doc.add_paragraph('✓ Eliminate manual calculation errors')
doc.add_paragraph('✓ Ensure accurate ingredient consumption')
doc.add_paragraph('✓ Display ingredient usage breakdown for each production')

doc.add_heading('Solution 11: Recipe Standardization', 3)
doc.add_paragraph('✓ Create standardized recipes with exact ingredient quantities')
doc.add_paragraph('✓ Link recipes to finished goods products')
doc.add_paragraph('✓ Ensure consistent product quality across all batches')
doc.add_paragraph('✓ Enable cost control through standardized ingredient usage')

doc.add_heading('Solution 12: Loss Factor Accounting', 3)
doc.add_paragraph('✓ Include loss factor percentage for each ingredient in recipes')
doc.add_paragraph('✓ Automatically adjust ingredient quantities for spillage, trimming, evaporation')
doc.add_paragraph('✓ Provide accurate ingredient consumption calculations')
doc.add_paragraph('✓ Prevent unexpected stock depletion')

doc.add_heading('Solution 13: Production History and Tracking', 3)
doc.add_paragraph('✓ Record all production activities with date, time, and user')
doc.add_paragraph('✓ Track which recipe was used and what quantity was produced')
doc.add_paragraph('✓ Show ingredients consumed for each production')
doc.add_paragraph('✓ Provide daily production summary on staff dashboard')
doc.add_paragraph('✓ Enable production efficiency analysis and quality tracing')

doc.add_heading('3.4 Solutions to Supplier and Purchase Order Problems', 2)

doc.add_heading('Solution 14: Digital Purchase Order System', 3)
doc.add_paragraph('✓ Create purchase orders digitally with auto-generated order numbers (PO-YYYYMMDD-XXXX)')
doc.add_paragraph('✓ Eliminate phone calls, emails, and paper forms')
doc.add_paragraph('✓ Ensure accurate order details')
doc.add_paragraph('✓ Maintain complete order history')

doc.add_heading('Solution 15: Supplier Portal for Collaboration', 3)
doc.add_paragraph('✓ Provide dedicated supplier login portal')
doc.add_paragraph('✓ Allow suppliers to view assigned purchase orders')
doc.add_paragraph('✓ Enable suppliers to approve orders and set pricing')
doc.add_paragraph('✓ Allow suppliers to update order status (Shipped)')
doc.add_paragraph('✓ Eliminate communication delays and errors')

doc.add_heading('Solution 16: Order Status Workflow Tracking', 3)
doc.add_paragraph('✓ Implement clear order status workflow: Draft → Pending → Approved → Shipped → Received')
doc.add_paragraph('✓ Display current status for all orders')
doc.add_paragraph('✓ Prevent duplicate orders through status visibility')
doc.add_paragraph('✓ Track expected delivery dates')

doc.add_heading('Solution 17: Automated Receiving Process', 3)
doc.add_paragraph('✓ Scan QR code to instantly identify order')
doc.add_paragraph('✓ Automatically create stock lots for all received items')
doc.add_paragraph('✓ Populate lot numbers, quantities, unit costs, and expiration dates')
doc.add_paragraph('✓ Eliminate manual data entry errors')
doc.add_paragraph('✓ Save time during receiving process')

doc.add_heading('Solution 18: QR Code Tracking System', 3)
doc.add_paragraph('✓ Generate unique QR code for each purchase order')
doc.add_paragraph('✓ Enable quick order identification during receiving')
doc.add_paragraph('✓ Match physical deliveries to digital records accurately')
doc.add_paragraph('✓ Reduce receiving errors and inventory discrepancies')

doc.add_heading('3.5 Solutions to Employee Attendance Problems', 2)

doc.add_heading('Solution 19: Digital Attendance System', 3)
doc.add_paragraph('✓ Implement digital clock in/out system')
doc.add_paragraph('✓ Prevent buddy punching through user authentication')
doc.add_paragraph('✓ Automatically record attendance with timestamps')
doc.add_paragraph('✓ Eliminate manual attendance compilation')

doc.add_heading('Solution 20: AM/PM Shift Tracking', 3)
doc.add_paragraph('✓ Support separate clock in/out for AM and PM shifts')
doc.add_paragraph('✓ Prevent duplicate clock-ins for same shift')
doc.add_paragraph('✓ Track partial attendance (AM only or PM only)')
doc.add_paragraph('✓ Simplify payroll calculation for shift-based work')

doc.add_heading('Solution 21: Attendance History and Calendar View', 3)
doc.add_paragraph('✓ Provide monthly calendar view with attendance status indicators')
doc.add_paragraph('✓ Show complete, partial, and missing attendance')
doc.add_paragraph('✓ Enable employees to view their own attendance history')
doc.add_paragraph('✓ Allow admins to review all staff attendance records')
doc.add_paragraph('✓ Support performance reviews and scheduling decisions')

doc.add_heading('3.6 Solutions to Reporting and Decision-Making Problems', 2)

doc.add_heading('Solution 22: Real-Time Dashboards', 3)
doc.add_paragraph('✓ Provide instant visibility into key metrics:')
doc.add_paragraph('  - Total products count')
doc.add_paragraph('  - Total inventory value')
doc.add_paragraph('  - Low stock and out-of-stock counts')
doc.add_paragraph('  - Finished goods statistics')
doc.add_paragraph('  - Recent items with stock status')
doc.add_paragraph('✓ Update metrics in real-time with every transaction')
doc.add_paragraph('✓ Enable data-driven decision making')

doc.add_heading('Solution 23: Trend Analysis and Visualization', 3)
doc.add_paragraph('✓ Display 6-month stock value trend chart')
doc.add_paragraph('✓ Show value change percentage compared to previous month')
doc.add_paragraph('✓ Enable identification of inventory patterns')
doc.add_paragraph('✓ Support strategic planning and forecasting')

doc.add_heading('Solution 24: Automated Report Generation', 3)
doc.add_paragraph('✓ Generate stock reports automatically (current levels, low stock, out of stock)')
doc.add_paragraph('✓ Provide expiration tracker report')
doc.add_paragraph('✓ Create production summaries')
doc.add_paragraph('✓ Offer searchable and filterable audit logs')
doc.add_paragraph('✓ Eliminate manual report compilation')

doc.add_page_break()

print("Creating Scope section...")
doc.save('Inventory_System_Documentation_Updated.docx')
print("Part 1 completed!")
