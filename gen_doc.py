from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Title Page
title = doc.add_heading('Inventory Management System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Complete System Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(16)

date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc = [
    '1. Introduction',
    '2. Standard Operating Procedures (SOP)',
    '3. Objectives',
    '4. Scope',
    '5. Limitations',
    '6. System Features'
]
for item in toc:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. INTRODUCTION
doc.add_heading('1. INTRODUCTION', 1)

doc.add_paragraph(
    'The Inventory Management System is a comprehensive Django-based web application designed to streamline '
    'inventory operations, user management, attendance tracking, and supplier collaboration. This enterprise-grade '
    'solution provides role-based access control with four distinct user types: Super Admin, Admin, Staff, and Supplier.'
)

doc.add_paragraph(
    'Built with security as the top priority, the system implements multiple layers of protection including SQL '
    'injection prevention, CSRF protection, XSS protection, and comprehensive audit logging. The application supports '
    'multi-user environments with granular permission management and complete activity tracking for compliance and security.'
)

doc.add_paragraph(
    'The system is designed for businesses that require precise inventory control, production tracking, supplier '
    'management, and employee attendance monitoring. It features automated workflows, real-time stock tracking, '
    'expiration management, and QR code-based purchase order processing.'
)

doc.add_heading('1.1 Technology Stack', 2)
doc.add_paragraph('• Backend Framework: Django 5.1.3')
doc.add_paragraph('• Database: MySQL 5.7+')
doc.add_paragraph('• Programming Language: Python 3.8+')
doc.add_paragraph('• Timezone: Asia/Manila')
doc.add_paragraph('• Security: Django ORM, CSRF Protection, HSTS, XSS Protection')

doc.add_page_break()

# 2. STANDARD OPERATING PROCEDURES
doc.add_heading('2. STANDARD OPERATING PROCEDURES (SOP)', 1)

doc.add_heading('2.1 User Management SOP', 2)
doc.add_paragraph('Step 1: Super Admin creates user accounts with appropriate roles (Super Admin, Admin, Staff, or Supplier)')
doc.add_paragraph('Step 2: Admin assigns granular permissions to staff users based on job responsibilities')
doc.add_paragraph('Step 3: Users login with their credentials through the unified login system')
doc.add_paragraph('Step 4: System automatically logs all user activities with IP address and timestamp')
doc.add_paragraph('Step 5: Admin monitors user activities through audit logs')
doc.add_paragraph('Step 6: Deactivate or delete users when necessary (Super Admin only for deletion)')

doc.add_heading('2.2 Inventory Management SOP', 2)
doc.add_paragraph('Step 1: Create items with auto-generated codes (Format: YYYYMM0001)')
doc.add_paragraph('Step 2: Set item category, unit, reorder level, and perishable status')
doc.add_paragraph('Step 3: Receive stock and create lots with lot numbers, quantities, and expiration dates')
doc.add_paragraph('Step 4: Monitor stock levels through dashboard alerts')
doc.add_paragraph('Step 5: Track expiration dates using the expiration tracker')
doc.add_paragraph('Step 6: Consume stock using FIFO (First In First Out) method')
doc.add_paragraph('Step 7: Generate reorder alerts when stock falls below reorder level')
doc.add_paragraph('Step 8: Perform stock adjustments when necessary with proper documentation')

doc.add_heading('2.3 Production SOP', 2)
doc.add_paragraph('Step 1: Create recipes with product selection and yield quantities')
doc.add_paragraph('Step 2: Add recipe ingredients with quantities, units, and loss factors')
doc.add_paragraph('Step 3: Initiate production by selecting recipe and entering production quantity')
doc.add_paragraph('Step 4: System automatically consumes ingredients using FIFO method')
doc.add_paragraph('Step 5: System creates finished goods stock lot with production details')
doc.add_paragraph('Step 6: Record production in stock movement history')
doc.add_paragraph('Step 7: Review production summary on staff dashboard')

doc.add_heading('2.4 Purchase Order SOP', 2)
doc.add_paragraph('Step 1: Create purchase order with supplier and items')
doc.add_paragraph('Step 2: Add items with quantities to the order')
doc.add_paragraph('Step 3: Submit order to supplier (system generates unique QR code)')
doc.add_paragraph('Step 4: Supplier logs in and reviews pending orders')
doc.add_paragraph('Step 5: Supplier approves order with unit pricing and expected delivery date')
doc.add_paragraph('Step 6: Supplier marks order as shipped when dispatched')
doc.add_paragraph('Step 7: Scan QR code upon delivery to receive order')
doc.add_paragraph('Step 8: System automatically creates stock lots for received items')
doc.add_paragraph('Step 9: Verify received quantities match ordered quantities')

doc.add_heading('2.5 Attendance Tracking SOP', 2)
doc.add_paragraph('Step 1: Staff member clocks in for AM shift at start of morning')
doc.add_paragraph('Step 2: Staff member clocks out for AM shift before lunch')
doc.add_paragraph('Step 3: Staff member clocks in for PM shift after lunch')
doc.add_paragraph('Step 4: Staff member clocks out for PM shift at end of day')
doc.add_paragraph('Step 5: System prevents duplicate clock-ins for same shift')
doc.add_paragraph('Step 6: Admin reviews attendance records and generates reports')
doc.add_paragraph('Step 7: View monthly calendar with attendance status indicators')

doc.add_heading('2.6 Supplier Portal SOP', 2)
doc.add_paragraph('Step 1: Admin creates supplier master data')
doc.add_paragraph('Step 2: Admin creates supplier user account and links to supplier')
doc.add_paragraph('Step 3: Supplier logs in through unified login')
doc.add_paragraph('Step 4: Supplier views assigned purchase orders')
doc.add_paragraph('Step 5: Supplier approves orders with pricing')
doc.add_paragraph('Step 6: Supplier updates order status to shipped')

doc.add_page_break()

# 3. OBJECTIVES
doc.add_heading('3. OBJECTIVES', 1)

doc.add_heading('3.1 Primary Objectives', 2)
doc.add_paragraph(
    '• Streamline Inventory Operations: Provide real-time tracking of stock levels, movements, and expiration dates '
    'to minimize waste and optimize inventory levels.'
)
doc.add_paragraph(
    '• Enhance Security: Implement comprehensive security measures including SQL injection protection, CSRF protection, '
    'audit logging, and role-based access control to protect sensitive business data.'
)
doc.add_paragraph(
    '• Improve Production Efficiency: Automate ingredient consumption and finished goods creation through recipe-based '
    'production workflows.'
)
doc.add_paragraph(
    '• Facilitate Supplier Collaboration: Enable seamless communication and order processing with suppliers through '
    'dedicated portal and QR code tracking.'
)
doc.add_paragraph(
    '• Monitor Employee Attendance: Track staff attendance with AM/PM shift support and calendar views for better '
    'workforce management.'
)

doc.add_heading('3.2 Secondary Objectives', 2)
doc.add_paragraph('• Reduce Manual Data Entry: Auto-generate codes, usernames, and order numbers')
doc.add_paragraph('• Minimize Stock Waste: Track perishable items and provide expiration alerts')
doc.add_paragraph('• Ensure Compliance: Maintain complete audit trail of all system activities')
doc.add_paragraph('• Support Decision Making: Provide dashboards with key metrics and trends')
doc.add_paragraph('• Enable Mobile Access: Support access from phones and tablets on local network')

doc.add_page_break()

# 4. SCOPE
doc.add_heading('4. SCOPE', 1)

doc.add_heading('4.1 Functional Scope', 2)

doc.add_heading('User Management', 3)
doc.add_paragraph('• User account creation, modification, and deletion')
doc.add_paragraph('• Role assignment (Super Admin, Admin, Staff, Supplier)')
doc.add_paragraph('• Granular permission management (10 permission types)')
doc.add_paragraph('• User relationship tracking (manager, subordinate, colleague)')
doc.add_paragraph('• Account activation/deactivation')

doc.add_heading('Inventory Management', 3)
doc.add_paragraph('• Item master data management (4 categories: Ingredient, Finished Good, Packaging, Equipment)')
doc.add_paragraph('• Stock lot tracking with lot numbers and expiration dates')
doc.add_paragraph('• Stock movement recording (6 types: Receive, Consume, Produce, Adjust, Transfer, Spoilage)')
doc.add_paragraph('• Real-time stock level calculations')
doc.add_paragraph('• Low stock and out-of-stock alerts')
doc.add_paragraph('• Expiration tracking and alerts')
doc.add_paragraph('• FIFO (First In First Out) stock consumption')

doc.add_heading('Production Management', 3)
doc.add_paragraph('• Recipe creation with ingredients and yield quantities')
doc.add_paragraph('• Loss factor calculation for ingredients')
doc.add_paragraph('• Automated ingredient consumption during production')
doc.add_paragraph('• Finished goods stock lot creation')
doc.add_paragraph('• Production history and tracking')

doc.add_heading('Purchase Order Management', 3)
doc.add_paragraph('• Purchase order creation and management')
doc.add_paragraph('• QR code generation for order tracking')
doc.add_paragraph('• Order workflow (Draft → Pending → Approved → Shipped → Received)')
doc.add_paragraph('• Supplier approval with pricing')
doc.add_paragraph('• QR code scanning for receiving')
doc.add_paragraph('• Automatic stock lot creation upon receipt')

doc.add_heading('Supplier Portal', 3)
doc.add_paragraph('• Supplier account management')
doc.add_paragraph('• Order viewing and approval')
doc.add_paragraph('• Pricing input during approval')
doc.add_paragraph('• Order status updates')
doc.add_paragraph('• Expected delivery date setting')

doc.add_heading('Attendance Tracking', 3)
doc.add_paragraph('• AM/PM shift clock in/out')
doc.add_paragraph('• Daily attendance records')
doc.add_paragraph('• Monthly calendar view')
doc.add_paragraph('• Attendance history')
doc.add_paragraph('• Admin attendance overview')

doc.add_heading('Reporting & Analytics', 3)
doc.add_paragraph('• Dashboard with key metrics')
doc.add_paragraph('• Stock value trends (6-month chart)')
doc.add_paragraph('• Low stock reports')
doc.add_paragraph('• Expiration tracker')
doc.add_paragraph('• Production summaries')
doc.add_paragraph('• Audit logs with filtering')

doc.add_heading('4.2 Technical Scope', 2)
doc.add_paragraph('• Web-based application accessible via browser')
doc.add_paragraph('• MySQL database backend')
doc.add_paragraph('• Django framework (Python)')
doc.add_paragraph('• Responsive design for desktop and mobile')
doc.add_paragraph('• Local network access support')
doc.add_paragraph('• Asia/Manila timezone support')
doc.add_paragraph('• UUID-based primary keys')
doc.add_paragraph('• RESTful API endpoints for dynamic data')

doc.add_heading('4.3 User Scope', 2)
doc.add_paragraph('• Super Admin: Full system access and control')
doc.add_paragraph('• Admin: User management and staff permission assignment')
doc.add_paragraph('• Staff: Attendance tracking and production operations')
doc.add_paragraph('• Supplier: Purchase order management through dedicated portal')

doc.add_page_break()

# 5. LIMITATIONS
doc.add_heading('5. LIMITATIONS', 1)

doc.add_heading('5.1 Technical Limitations', 2)
doc.add_paragraph(
    '• Single Database: The system currently supports only MySQL database. Migration to other databases '
    'would require configuration changes.'
)
doc.add_paragraph(
    '• Local Network Deployment: Designed for local network deployment. Cloud deployment would require '
    'additional security configurations and HTTPS setup.'
)
doc.add_paragraph(
    '• No Multi-location Support: The system does not currently support multiple warehouse locations or '
    'inter-location transfers.'
)
doc.add_paragraph(
    '• Limited Reporting: Advanced reporting features like custom report builders and data export to Excel '
    'are not implemented.'
)
doc.add_paragraph(
    '• No Email Notifications: The system does not send automated email notifications for low stock, '
    'expiring items, or order updates.'
)

doc.add_heading('5.2 Functional Limitations', 2)
doc.add_paragraph(
    '• No Barcode Scanning: While QR codes are supported for purchase orders, general barcode scanning '
    'for items is not implemented.'
)
doc.add_paragraph(
    '• No Multi-currency Support: All pricing is in single currency. Multi-currency support is not available.'
)
doc.add_paragraph(
    '• No Integration with Accounting: The system does not integrate with external accounting software.'
)
doc.add_paragraph(
    '• Limited Batch Operations: Bulk operations are limited. Mass updates require individual processing.'
)
doc.add_paragraph(
    '• No Mobile App: Access is through web browser only. Native mobile applications are not available.'
)

doc.add_heading('5.3 Business Limitations', 2)
doc.add_paragraph(
    '• Single Organization: The system is designed for single organization use. Multi-tenant support '
    'is not available.'
)
doc.add_paragraph(
    '• No Customer Management: The system focuses on inventory and suppliers. Customer relationship '
    'management features are not included.'
)
doc.add_paragraph(
    '• No Sales Module: While production is tracked, sales order management is not implemented.'
)
doc.add_paragraph(
    '• Limited Shift Management: Only AM/PM shifts are supported. Complex shift patterns are not available.'
)

doc.add_heading('5.4 Security Limitations', 2)
doc.add_paragraph(
    '• No Two-Factor Authentication: The system uses username/password authentication only.'
)
doc.add_paragraph(
    '• No Password Recovery: Password reset functionality requires admin intervention.'
)
doc.add_paragraph(
    '• Limited Session Management: Advanced session controls like concurrent login prevention are not implemented.'
)

doc.add_heading('5.5 Performance Limitations', 2)
doc.add_paragraph(
    '• Not Optimized for Large Scale: The system is designed for small to medium businesses. '
    'Performance with millions of records has not been tested.'
)
doc.add_paragraph(
    '• No Caching: Advanced caching mechanisms for improved performance are not implemented.'
)
doc.add_paragraph(
    '• Sequential Processing: Batch operations are processed sequentially, not in parallel.'
)

doc.add_page_break()

print("Creating features section...")
doc.save('Inventory_System_Documentation.docx')
print("Documentation created successfully!")
