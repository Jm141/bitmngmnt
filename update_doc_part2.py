from docx import Document

# Open existing document
doc = Document('Inventory_System_Documentation_Updated.docx')

# 4. SCOPE
doc.add_heading('4. SCOPE', 1)

doc.add_heading('4.1 Functional Scope', 2)

doc.add_heading('User Management', 3)
doc.add_paragraph('• User account creation, modification, and deletion')
doc.add_paragraph('• Role assignment (Super Admin, Admin, Staff, Supplier)')
doc.add_paragraph('• Granular permission management (10 permission types)')
doc.add_paragraph('• User relationship tracking (manager, subordinate, colleague)')
doc.add_paragraph('• Account activation/deactivation')

doc.add_heading('Inventory Management', 3)
doc.add_paragraph('• Item master data management (4 categories: Ingredient, Finished Good, Packaging, Equipment)')
doc.add_paragraph('• Stock lot tracking with lot numbers and expiration dates')
doc.add_paragraph('• Stock movement recording (6 types: Receive, Consume, Produce, Adjust, Transfer, Spoilage)')
doc.add_paragraph('• Real-time stock level calculations')
doc.add_paragraph('• Low stock and out-of-stock alerts')
doc.add_paragraph('• Expiration tracking and alerts')
doc.add_paragraph('• FIFO (First In First Out) stock consumption')

doc.add_heading('Production Management', 3)
doc.add_paragraph('• Recipe creation with ingredients and yield quantities')
doc.add_paragraph('• Loss factor calculation for ingredients')
doc.add_paragraph('• Automated ingredient consumption during production')
doc.add_paragraph('• Finished goods stock lot creation')
doc.add_paragraph('• Production history and tracking')

doc.add_heading('Purchase Order Management', 3)
doc.add_paragraph('• Purchase order creation and management')
doc.add_paragraph('• QR code generation for order tracking')
doc.add_paragraph('• Order workflow (Draft → Pending → Approved → Shipped → Received)')
doc.add_paragraph('• Supplier approval with pricing')
doc.add_paragraph('• QR code scanning for receiving')
doc.add_paragraph('• Automatic stock lot creation upon receipt')

doc.add_heading('Supplier Portal', 3)
doc.add_paragraph('• Supplier account management')
doc.add_paragraph('• Order viewing and approval')
doc.add_paragraph('• Pricing input during approval')
doc.add_paragraph('• Order status updates')
doc.add_paragraph('• Expected delivery date setting')

doc.add_heading('Attendance Tracking', 3)
doc.add_paragraph('• AM/PM shift clock in/out')
doc.add_paragraph('• Daily attendance records')
doc.add_paragraph('• Monthly calendar view')
doc.add_paragraph('• Attendance history')
doc.add_paragraph('• Admin attendance overview')

doc.add_heading('Reporting & Analytics', 3)
doc.add_paragraph('• Dashboard with key metrics')
doc.add_paragraph('• Stock value trends (6-month chart)')
doc.add_paragraph('• Low stock reports')
doc.add_paragraph('• Expiration tracker')
doc.add_paragraph('• Production summaries')
doc.add_paragraph('• Audit logs with filtering')

doc.add_heading('4.2 Technical Scope', 2)
doc.add_paragraph('• Web-based application accessible via browser')
doc.add_paragraph('• MySQL database backend')
doc.add_paragraph('• Django framework (Python)')
doc.add_paragraph('• Responsive design for desktop and mobile')
doc.add_paragraph('• Local network access support')
doc.add_paragraph('• Asia/Manila timezone support')
doc.add_paragraph('• UUID-based primary keys')
doc.add_paragraph('• RESTful API endpoints for dynamic data')

doc.add_heading('4.3 User Scope', 2)
doc.add_paragraph('• Super Admin: Full system access and control')
doc.add_paragraph('• Admin: User management and staff permission assignment')
doc.add_paragraph('• Staff: Attendance tracking and production operations')
doc.add_paragraph('• Supplier: Purchase order management through dedicated portal')

doc.add_page_break()

# 5. LIMITATIONS
doc.add_heading('5. LIMITATIONS', 1)

doc.add_heading('5.1 Technical Limitations', 2)
doc.add_paragraph(
    '• Single Database: The system currently supports only MySQL database. Migration to other databases '
    'would require configuration changes.'
)
doc.add_paragraph(
    '• Local Network Deployment: Designed for local network deployment. Cloud deployment would require '
    'additional security configurations and HTTPS setup.'
)
doc.add_paragraph(
    '• No Multi-location Support: The system does not currently support multiple warehouse locations or '
    'inter-location transfers.'
)
doc.add_paragraph(
    '• Limited Reporting: Advanced reporting features like custom report builders and data export to Excel '
    'are not implemented.'
)
doc.add_paragraph(
    '• No Email Notifications: The system does not send automated email notifications for low stock, '
    'expiring items, or order updates.'
)

doc.add_heading('5.2 Functional Limitations', 2)
doc.add_paragraph(
    '• No Barcode Scanning: While QR codes are supported for purchase orders, general barcode scanning '
    'for items is not implemented.'
)
doc.add_paragraph(
    '• No Multi-currency Support: All pricing is in single currency. Multi-currency support is not available.'
)
doc.add_paragraph(
    '• No Integration with Accounting: The system does not integrate with external accounting software.'
)
doc.add_paragraph(
    '• Limited Batch Operations: Bulk operations are limited. Mass updates require individual processing.'
)
doc.add_paragraph(
    '• No Mobile App: Access is through web browser only. Native mobile applications are not available.'
)

doc.add_heading('5.3 Business Limitations', 2)
doc.add_paragraph(
    '• Single Organization: The system is designed for single organization use. Multi-tenant support '
    'is not available.'
)
doc.add_paragraph(
    '• No Customer Management: The system focuses on inventory and suppliers. Customer relationship '
    'management features are not included.'
)
doc.add_paragraph(
    '• No Sales Module: While production is tracked, sales order management is not implemented.'
)
doc.add_paragraph(
    '• Limited Shift Management: Only AM/PM shifts are supported. Complex shift patterns are not available.'
)

doc.add_heading('5.4 Security Limitations', 2)
doc.add_paragraph(
    '• No Two-Factor Authentication: The system uses username/password authentication only.'
)
doc.add_paragraph(
    '• No Password Recovery: Password reset functionality requires admin intervention.'
)
doc.add_paragraph(
    '• Limited Session Management: Advanced session controls like concurrent login prevention are not implemented.'
)

doc.add_heading('5.5 Performance Limitations', 2)
doc.add_paragraph(
    '• Not Optimized for Large Scale: The system is designed for small to medium businesses. '
    'Performance with millions of records has not been tested.'
)
doc.add_paragraph(
    '• No Caching: Advanced caching mechanisms for improved performance are not implemented.'
)
doc.add_paragraph(
    '• Sequential Processing: Batch operations are processed sequentially, not in parallel.'
)

doc.add_page_break()

# 6. SYSTEM FEATURES (Summary)
doc.add_heading('6. SYSTEM FEATURES', 1)

doc.add_paragraph(
    'The following features have been implemented to solve the problems identified in Section 2 '
    'and achieve the objectives outlined in Section 3.'
)

doc.add_heading('6.1 Security & Authentication Features', 2)
doc.add_paragraph('• Unified login system for all user types')
doc.add_paragraph('• Role-based authentication (Super Admin, Admin, Staff, Supplier)')
doc.add_paragraph('• SQL Injection Protection via Django ORM')
doc.add_paragraph('• CSRF Protection on all forms')
doc.add_paragraph('• XSS Protection through input validation')
doc.add_paragraph('• Secure headers (HSTS, X-Frame-Options)')
doc.add_paragraph('• Comprehensive audit logging with IP tracking')
doc.add_paragraph('• Password validation and encryption')

doc.add_heading('6.2 User Management Features', 2)
doc.add_paragraph('• User CRUD operations (Create, Read, Update, Delete)')
doc.add_paragraph('• Auto-generated usernames from first/last name')
doc.add_paragraph('• Role assignment and hierarchy enforcement')
doc.add_paragraph('• Granular permission system (10 permission types)')
doc.add_paragraph('• Bulk permission updates')
doc.add_paragraph('• User relationship tracking')
doc.add_paragraph('• Account activation/deactivation')

doc.add_heading('6.3 Inventory Management Features', 2)
doc.add_paragraph('• Auto-generated item codes (YYYYMM0001 format)')
doc.add_paragraph('• 4 item categories (Ingredient, Finished Good, Packaging, Equipment)')
doc.add_paragraph('• 8 unit types (pcs, kg, g, L, mL, pack, box, dozen)')
doc.add_paragraph('• Stock lot tracking with lot numbers')
doc.add_paragraph('• Expiration date management')
doc.add_paragraph('• FIFO (First In First Out) consumption')
doc.add_paragraph('• Real-time stock level calculations')
doc.add_paragraph('• Low stock and out-of-stock alerts')
doc.add_paragraph('• Expiration tracker (7-day alert threshold)')
doc.add_paragraph('• 6 stock movement types (Receive, Consume, Produce, Adjust, Transfer, Spoilage)')
doc.add_paragraph('• Complete movement history with user tracking')

doc.add_heading('6.4 Production Features', 2)
doc.add_paragraph('• Recipe creation with standardized ingredients')
doc.add_paragraph('• Yield quantity specification')
doc.add_paragraph('• Loss factor calculation for ingredients')
doc.add_paragraph('• Automated ingredient consumption during production')
doc.add_paragraph('• Automatic finished goods stock lot creation')
doc.add_paragraph('• Production history tracking')
doc.add_paragraph('• Daily production summary on staff dashboard')
doc.add_paragraph('• Ingredient usage breakdown display')

doc.add_heading('6.5 Purchase Order Features', 2)
doc.add_paragraph('• Auto-generated order numbers (PO-YYYYMMDD-XXXX format)')
doc.add_paragraph('• Unique QR code generation for each order')
doc.add_paragraph('• Order status workflow (Draft → Pending → Approved → Shipped → Received)')
doc.add_paragraph('• Supplier portal for order viewing and approval')
doc.add_paragraph('• Pricing input by suppliers during approval')
doc.add_paragraph('• Expected delivery date setting')
doc.add_paragraph('• QR code scanning for receiving')
doc.add_paragraph('• Automatic stock lot creation upon receipt')
doc.add_paragraph('• Order cancellation capability')

doc.add_heading('6.6 Attendance/DTR Features', 2)
doc.add_paragraph('• AM/PM shift clock in/out')
doc.add_paragraph('• Manila timezone support')
doc.add_paragraph('• Single clock-in per shift enforcement')
doc.add_paragraph('• Daily attendance records')
doc.add_paragraph('• Monthly calendar view with status indicators')
doc.add_paragraph('• Complete/Partial/Missing attendance display')
doc.add_paragraph('• Attendance history')
doc.add_paragraph('• Admin attendance overview for all staff')

doc.add_heading('6.7 Dashboard & Reporting Features', 2)
doc.add_paragraph('• Real-time dashboard with key metrics')
doc.add_paragraph('• Total inventory value display')
doc.add_paragraph('• Low stock and out-of-stock counts')
doc.add_paragraph('• Finished goods statistics')
doc.add_paragraph('• 6-month stock value trend chart')
doc.add_paragraph('• Value change percentage')
doc.add_paragraph('• Recent activity logs')
doc.add_paragraph('• Stock reports (current levels, low stock, out of stock)')
doc.add_paragraph('• Expiration tracker report')
doc.add_paragraph('• Production summaries')
doc.add_paragraph('• Searchable and filterable audit logs')

doc.add_heading('6.8 Additional Features', 2)
doc.add_paragraph('• Search and filtering (users, items, dates)')
doc.add_paragraph('• Pagination support')
doc.add_paragraph('• RESTful API endpoints for dynamic data')
doc.add_paragraph('• Input sanitization and validation')
doc.add_paragraph('• Responsive design for mobile access')
doc.add_paragraph('• Local network access support')
doc.add_paragraph('• CSRF-protected forms')
doc.add_paragraph('• Error and success messaging')

doc.add_page_break()

# System Statistics
doc.add_heading('7. SYSTEM STATISTICS', 1)
doc.add_paragraph('• Total Database Models: 14')
doc.add_paragraph('• Total View Functions: 40+')
doc.add_paragraph('• Total URL Patterns: 40+')
doc.add_paragraph('• User Roles: 4 (Super Admin, Admin, Staff, Supplier)')
doc.add_paragraph('• Permission Types: 10')
doc.add_paragraph('• Stock Movement Types: 6')
doc.add_paragraph('• Item Categories: 4')
doc.add_paragraph('• Unit Types: 8')
doc.add_paragraph('• Order Status Types: 6')

doc.add_page_break()

# Conclusion
doc.add_heading('8. CONCLUSION', 1)

doc.add_paragraph(
    'The Inventory Management System comprehensively addresses the 24 problems identified in the Statement of the Problem '
    'through targeted, automated solutions. By implementing real-time tracking, role-based security, automated workflows, '
    'and comprehensive reporting, the system transforms traditional manual inventory management into an efficient, '
    'secure, and data-driven operation.'
)

doc.add_paragraph(
    'Key achievements include:'
)
doc.add_paragraph('• Elimination of manual tracking errors through automation')
doc.add_paragraph('• Reduction of waste through expiration management and FIFO')
doc.add_paragraph('• Enhanced security through multi-layer protection and audit logging')
doc.add_paragraph('• Improved production efficiency through recipe standardization')
doc.add_paragraph('• Streamlined supplier collaboration through dedicated portal')
doc.add_paragraph('• Accurate attendance tracking with shift support')
doc.add_paragraph('• Real-time visibility into business operations through dashboards')

doc.add_paragraph(
    'The system is designed for scalability and can be extended with additional features such as email notifications, '
    'barcode scanning, multi-location support, and accounting integration as business needs evolve.'
)

# Save document
doc.save('Inventory_System_Documentation_Updated.docx')
print("Complete documentation created successfully!")
print("File: Inventory_System_Documentation_Updated.docx")
print("\nDocument structure:")
print("1. Introduction")
print("2. Statement of the Problem (SOP) - 24 problems identified")
print("3. Objectives (Solutions) - Solutions to all 24 problems")
print("4. Scope")
print("5. Limitations")
print("6. System Features")
print("7. System Statistics")
print("8. Conclusion")
